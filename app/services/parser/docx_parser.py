import logging
from docx import Document
from langsmith import traceable

from app.services.parser.base import BaseParser
from app.schemas.document import (
    ParsedDocument,
    Page,
    TextLine,
)

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):

    @traceable(
        name="DOCX Parser",
        run_type="parser",
    )
    def parse(self, file_path: str) -> ParsedDocument:
        logger.info("Parsing DOCX: %s", file_path)

        doc = Document(file_path)

        lines = []

        for para in doc.paragraphs:

            text = para.text.strip()

            if not text:
                continue

            # Default values
            font = None
            font_size = None
            is_bold = False

            # Extract formatting from runs
            if para.runs:
                first_run = para.runs[0]

                if first_run.font.name:
                    font = first_run.font.name

                if first_run.font.size:
                    # Convert EMU to points
                    font_size = first_run.font.size.pt

                is_bold = any(
                    run.bold for run in para.runs
                )

            lines.append(
                TextLine(
                    text=text,
                    page=1,
                    bbox=None,      # DOCX has no coordinates
                    font=font,
                    font_size=font_size,
                    is_bold=is_bold,
                )
            )

        page = Page(
            page_number=1,
            text="\n".join(line.text for line in lines),
            lines=lines,
        )

        logger.info(
            "Parsed %d paragraphs from %s",
            len(lines),
            file_path,
        )

        return ParsedDocument(
            filename=file_path,
            total_pages=1,
            pages=[page],
        )